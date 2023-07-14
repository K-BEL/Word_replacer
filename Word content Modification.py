from docx import Document
import pandas as pd
import tkinter as tk
import tkinter.ttk as ttk
from tkinter.filedialog import askopenfile, askopenfilename
from tkinter.messagebox import showinfo

win = tk.Tk()
win.title("Modification du document word")

my_dic = {}


def openfile():
    file = askopenfilename(filetypes=[('Word Files', '*.docx')])
    print(file)
    document = Document(file)
    showinfo("Done", "File Successfully selected")
    print(my_dic)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if 'Plle' in paragraph.text:
                        m = ''
                        for j in paragraph.text:
                            if j.isdigit():
                                m += str(j)
                        try:
                            paragraph.text = 'REQ' + str(my_dic['REQ'][my_dic['Plle'][int(m)]])
                        except:
                            pass
    c = ''
    for p in document.paragraphs:
        if 'parcelle n°' in p.text.lower():
            for i in p.text:
                if i.isdigit():
                    c += str(i)
        if 'Réquisition n° :' in p.text or 'Réquisition d’immatriculationN°' in p.text:
            if c == '':
                pp=p
                pp.text = ''
            else:
                p.text = 'REQ' + str(my_dic['REQ'][my_dic['Plle'][int(c)]])
            c = ''
    document.save('new word file.docx')


def openfile1():
    global my_dic
    file = askopenfilename(filetypes=[('excel Files', '*.xlsx')])
    print(file)
    my_dic = pd.read_excel(file).to_dict()
    my_dic['Plle'] = {v: k for k, v in my_dic['Plle'].items()}
    print(my_dic)
    showinfo("Done", "File Successfully selected")


label = tk.Label(win, text='Choose excel File: ')
label.grid(row=0, column=0, padx=5, pady=5)

label = tk.Label(win, text='Choose word File: ')
label.grid(row=1, column=0, padx=5, pady=5)

button = ttk.Button(win, text='choisir fichier word ', width=30, command=openfile)
button.grid(row=1, column=1, padx=5, pady=5)

button = ttk.Button(win, text='choisir fichier excel', width=30, command=openfile1)
button.grid(row=0, column=1, padx=5, pady=5)

win.mainloop()
